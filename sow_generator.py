#!/usr/bin/env python3
"""
╔═══════════════════════════════════════════════════════════════╗
║          SOW Orchestrator — Python CLI                        ║
║          AI Portfolio · Daniel Mazzini                        ║
║                                                               ║
║  Autonomous 3-agent pipeline:                                 ║
║    Agent 1 · Discovery Analyst  — structures raw notes        ║
║    Agent 2 · Scope Validator    — autonomous go/no-go gate    ║
║    Agent 3 · SOW Drafter        — writes SOW if approved      ║
║                                                               ║
║  Providers: Claude (claude-haiku-4-5) · OpenAI (gpt-4o-mini) ║
╚═══════════════════════════════════════════════════════════════╝
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

# ─── DEPENDENCY CHECKS ────────────────────────────────────────────────────────

try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

try:
    import openai as openai_lib
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False

try:
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import Progress, SpinnerColumn, TextColumn, TimeElapsedColumn
    from rich.table import Table
    from rich.rule import Rule
    HAS_RICH = True
    console = Console()
except ImportError:
    HAS_RICH = False
    class Console:
        def print(self, *args, **kwargs): print(*args)
        def log(self, *args, **kwargs): print(*args)
        def input(self, prompt=""): return input(prompt)
    console = Console()

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ─── CONSTANTS ────────────────────────────────────────────────────────────────

VERSION        = "2.0.0"
CLAUDE_MODEL   = "claude-haiku-4-5-20251001"
OPENAI_MODEL   = "gpt-4o-mini"
GOLD           = "#C9A84C"
CHARCOAL       = "#2C2C2C"

ENGAGEMENT_TYPES = [
    "Private Cloud Deployment (VCF)",
    "Network Security Hardening (NSX)",
    "Telco NFV Transformation",
    "SASE / SD-WAN Migration",
    "Cloud Modernization Assessment",
    "NOC Automation Program",
    "Data Center Consolidation",
    "Hybrid Cloud Architecture",
]

ENGAGEMENT_SIZES = [
    "Small (under $150K)",
    "Medium ($150K–$500K)",
    "Large ($500K–$2M)",
    "Enterprise ($2M+)",
]

INDUSTRIES = [
    "Financial Services", "Healthcare", "Telco / CSP",
    "Manufacturing", "Technology", "Government",
    "Energy", "Retail", "Not specified",
]

# ─── JSON REPAIR ──────────────────────────────────────────────────────────────

def repair_json(raw: str) -> dict:
    """4-attempt JSON recovery chain — matches web tool logic."""
    first = raw.find('{')
    last  = raw.rfind('}')
    if first == -1 or last == -1:
        raise ValueError("No JSON object found in response")
    chunk = raw[first:last + 1]

    for transform in [
        lambda s: s,
        lambda s: re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', s),
        lambda s: re.sub(r'(?<!\\)\n', r'\\n', s),
        lambda s: re.sub(r'\r?\n|\t', ' ', re.sub(r'\s{2,}', ' ', s)),
    ]:
        try:
            return json.loads(transform(chunk))
        except json.JSONDecodeError:
            continue

    raise ValueError("All JSON repair attempts failed — please try again.")


def sanitise(text: str) -> str:
    """Remove characters that break JSON string values in prompts."""
    return (
        text.replace('\\', '/')
            .replace('"', "'")
            .replace('\r\n', ' | ')
            .replace('\n', ' | ')
            .replace('\t', ' ')
    )

# ─── PROVIDER ABSTRACTION ─────────────────────────────────────────────────────

def call_provider(provider: str, prompt: str, max_tokens: int,
                  claude_client=None, openai_client=None) -> str:
    """Route to the selected provider and return the raw text response."""
    if provider == "claude":
        if claude_client is None:
            raise RuntimeError("Claude client not initialised")
        msg = claude_client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text

    elif provider == "openai":
        if openai_client is None:
            raise RuntimeError("OpenAI client not initialised")
        resp = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}],
        )
        return resp.choices[0].message.content

    else:
        raise ValueError(f"Unknown provider: {provider}")

# ─── AGENT 1: DISCOVERY ANALYST ───────────────────────────────────────────────

def run_agent1(provider: str, notes: str, eng_type: str, eng_size: str,
               industry: str, claude_client=None, openai_client=None) -> dict:
    safe_notes = sanitise(notes)
    prompt = f"""You are a Senior Solutions Architect and Technical Program Manager with 20+ years structuring professional services engagements.

CRITICAL JSON RULES: Return ONLY a raw JSON object. No markdown, no backticks, no explanation. Every string value must be on a single line with no line breaks inside it. Use single quotes inside strings. Start with {{ end with }}.

ENGAGEMENT TYPE: {eng_type}
ENGAGEMENT SIZE: {eng_size}
CUSTOMER INDUSTRY: {industry}

RAW DISCOVERY NOTES:
{safe_notes}

Return exactly this JSON structure — all string values SHORT and on ONE LINE:
{{
  "engagement_objective": "one sentence describing success",
  "client_profile": "two sentences on customer and context",
  "in_scope": ["workstream 1", "workstream 2", "workstream 3"],
  "out_of_scope": ["excluded item 1", "excluded item 2"],
  "deliverables": [
    {{"name": "deliverable name", "description": "what it includes", "format": "document or workshop or implementation"}}
  ],
  "phases": [
    {{"phase": "phase name", "duration": "e.g. 4 weeks", "activities": ["activity 1", "activity 2"]}}
  ],
  "assumptions": ["assumption 1", "assumption 2", "assumption 3"],
  "customer_responsibilities": ["responsibility 1", "responsibility 2"],
  "dependencies": ["dependency 1", "dependency 2"],
  "success_criteria": ["criterion 1", "criterion 2"],
  "risks": [
    {{"risk": "risk title", "mitigation": "how to address it"}}
  ],
  "estimated_duration": "total duration e.g. 14 weeks",
  "team_structure": "recommended team composition",
  "open_items": ["open item 1", "open item 2"]
}}"""

    raw = call_provider(provider, prompt, 4000, claude_client, openai_client)
    if not raw:
        raise ValueError("Agent 1 returned an empty response")
    return repair_json(raw)

# ─── AGENT 2: SCOPE VALIDATOR ─────────────────────────────────────────────────

def run_agent2(provider: str, brief: dict, eng_type: str, eng_size: str,
               industry: str, claude_client=None, openai_client=None) -> dict:
    prompt = f"""You are a Senior QA Architect reviewing a professional services requirements brief before SOW drafting.

Return ONLY raw JSON — no markdown, no backticks. All string values on ONE LINE. Single quotes inside strings. Start with {{ end with }}.

ENGAGEMENT TYPE: {eng_type}
ENGAGEMENT SIZE: {eng_size}
CUSTOMER INDUSTRY: {industry}

BRIEF TO VALIDATE:
{json.dumps(brief)}

Check for: contradictions, missing critical items, ambiguities, unrealistic timelines, scope gaps.

Return exactly this JSON — all string values SHORT and on ONE LINE:
{{
  "decision": "approved",
  "validation_summary": "one sentence summary",
  "quality_score": 85,
  "issues": [
    {{"type": "Contradiction", "title": "short title", "description": "what the problem is", "recommendation": "what to fix"}}
  ],
  "warnings": [
    {{"type": "Risk", "title": "short title", "description": "what to watch for"}}
  ],
  "approved_items": ["solid item 1", "solid item 2", "solid item 3"]
}}

Decision: "rejected" if critical issues, "approved_with_warnings" if only warnings, "approved" if clean.
Maximum 4 issues, 4 warnings, 4 approved_items."""

    raw = call_provider(provider, prompt, 3000, claude_client, openai_client)
    if not raw:
        raise ValueError("Agent 2 returned an empty response")
    return repair_json(raw)

# ─── AGENT 3: SOW DRAFTER ─────────────────────────────────────────────────────

def run_agent3(provider: str, brief: dict, validation: dict,
               eng_type: str, eng_size: str, industry: str,
               claude_client=None, openai_client=None) -> str:
    warnings_ctx = ""
    if validation.get("warnings"):
        warnings_ctx = "\n\nVALIDATION NOTES (address in SOW):\n"
        warnings_ctx += "\n".join(
            f"- {w['title']}: {w['description']}" for w in validation["warnings"]
        )

    today     = datetime.today().strftime("%B %d, %Y")
    today_iso = datetime.today().strftime("%Y-%m-%d")

    prompt = f"""You are a Senior Principal Consultant with 20+ years writing Statements of Work for VMware, Dell EMC, Juniper Networks, and Nokia/Ericsson. This brief has been reviewed and approved by a Scope Validator. Write the final SOW with confidence.

Write a Statement of Work in clean Markdown. Use tables wherever possible. Be specific — reference actual technologies and customer context. Every section opens with 1-2 sentences of professional prose, then a table.

ENGAGEMENT TYPE: {eng_type}
ENGAGEMENT SIZE: {eng_size}
CUSTOMER INDUSTRY: {industry}

VALIDATED BRIEF:
{json.dumps(brief, indent=2)}{warnings_ctx}

Write exactly these 14 sections:

# {eng_type} — Statement of Work
**Version:** 1.0 | **Date:** {today} | **Status:** Draft for Review | **Prepared by:** Daniel Mazzini, Principal Architect

## 1. Engagement Overview
[2 sentences on customer context and why this engagement matters]
| Field | Detail |
|---|---|
[rows: Engagement Type, Industry, Size, Duration, Delivery Model, Engagement Lead]

## 2. Engagement Objective
[2 specific sentences referencing actual technologies and business outcome]

## 3. Scope of Work
[1 sentence]
| Workstream | Description | Key Activities | In Scope |
[one row per workstream]

## 4. Deliverables
[1 sentence]
| # | Deliverable | Description | Format | Qty | Est. Effort (hrs) |
[each deliverable with realistic hours]
| **TOTAL** | | | | | **XX hrs** |

## 5. Work Breakdown & Schedule
[1 sentence]
| Phase | Activity | Owner | Duration | Hours |
[grouped by phase with subtotals. Final row: TOTAL]

## 6. Milestones & Acceptance
[1 sentence]
| # | Milestone | Acceptance Criteria | Target Week | Owner |

## 7. Assumptions
[1 sentence]
| # | Assumption | Dependency | Risk if Invalid |

## 8. Customer Responsibilities
[1 sentence]
| # | Responsibility | Details | Required By | Owner |

## 9. Out of Scope
[1 sentence]
| # | Item | Rationale | How to Add |

## 10. Risk Register
[1 sentence]
| # | Risk | Probability | Impact | Mitigation | Owner |

## 11. Team Structure
[1 sentence]
| Role | Responsibilities | Allocation | Location |

## 12. Open Items
[1 sentence]
| # | Item | Decision Required | Owner | Due Date |

## 13. Commercial Terms
[1 sentence]
| Term | Details |
[rows: Engagement Model, Payment Schedule, Travel & Expenses, Change Order Process, Cancellation Notice, Warranty Period]

## 14. Document Control
| Version | Date | Author | Changes |
| 1.0 | {today_iso} | Daniel Mazzini, Principal Architect | Initial draft |

---
*This Statement of Work is a draft for review and is subject to legal review, customer negotiation, and final approval before execution. Generated by AI SOW Orchestrator · Daniel Mazzini · Principal Architect & Senior TPM.*"""

    raw = call_provider(provider, prompt, 8000, claude_client, openai_client)
    return raw or ""

# ─── RICH DISPLAY HELPERS ─────────────────────────────────────────────────────

def print_header(provider: str):
    provider_label = f"Claude ({CLAUDE_MODEL})" if provider == "claude" else f"OpenAI ({OPENAI_MODEL})"
    if not HAS_RICH:
        print(f"\n=== SOW Orchestrator — AI Portfolio ===")
        print(f"    Provider: {provider_label}")
        print(f"    Daniel Mazzini · 3-Agent Pipeline\n")
        return
    console.print()
    console.print(Panel.fit(
        f"[bold #C9A84C]SOW Orchestrator[/bold #C9A84C]  ·  [dim]AI Portfolio · Daniel Mazzini[/dim]\n"
        f"[dim]Provider:[/dim] [bold]{provider_label}[/bold]   "
        f"[dim]Agents:[/dim] Discovery Analyst · Scope Validator · SOW Drafter",
        border_style="#C9A84C",
        padding=(1, 4),
    ))
    console.print()


def print_agent_start(num: int, name: str, desc: str):
    if not HAS_RICH:
        print(f"\n[Agent {num}] {name}")
        print(f"  {desc}")
        return
    colors = {1: "#C9A84C", 2: "#FB923C", 3: "#34D399"}
    color  = colors.get(num, "white")
    console.print(Rule(f"[bold {color}]Agent {num} · {name}[/bold {color}]", style=color))
    console.print(f"  [dim]{desc}[/dim]")


def print_agent_done(num: int, summary: str):
    if not HAS_RICH:
        print(f"  ✓ {summary}")
        return
    console.print(f"  [green]✓[/green]  {summary}")


def print_agent_error(num: int, msg: str):
    if not HAS_RICH:
        print(f"  ✗ Error in Agent {num}: {msg}")
        return
    console.print(f"  [red]✗ Error in Agent {num}:[/red] {msg}")


def print_brief_summary(brief: dict):
    if not HAS_RICH:
        print(f"\n  Objective:    {brief.get('engagement_objective','')[:80]}")
        print(f"  Workstreams:  {len(brief.get('in_scope',[]))}")
        print(f"  Deliverables: {len(brief.get('deliverables',[]))}")
        print(f"  Phases:       {len(brief.get('phases',[]))}")
        print(f"  Risks:        {len(brief.get('risks',[]))}")
        print(f"  Duration:     {brief.get('estimated_duration','—')}")
        return

    t = Table(show_header=False, box=None, padding=(0, 2))
    t.add_column(style="dim", width=20)
    t.add_column()
    t.add_row("Objective",    brief.get("engagement_objective", "")[:90])
    t.add_row("Workstreams",  str(len(brief.get("in_scope", []))))
    t.add_row("Deliverables", str(len(brief.get("deliverables", []))))
    t.add_row("Phases",       str(len(brief.get("phases", []))))
    t.add_row("Risks",        str(len(brief.get("risks", []))))
    t.add_row("Duration",     brief.get("estimated_duration", "—"))
    t.add_row("Team",         brief.get("team_structure", "—"))
    console.print(Panel(t,
        title="[bold]Agent 1 — Requirements Brief[/bold]",
        border_style="#C9A84C", padding=(0, 1)))


def print_validation_result(v: dict):
    decision = v.get("decision", "approved")
    score    = v.get("quality_score", 0)
    issues   = v.get("issues", [])
    warnings = v.get("warnings", [])

    color_map = {"approved": "green", "approved_with_warnings": "yellow", "rejected": "red"}
    icon_map  = {"approved": "✓", "approved_with_warnings": "⚠", "rejected": "✗"}
    color     = color_map.get(decision, "white")
    icon      = icon_map.get(decision, "?")

    if not HAS_RICH:
        label = decision.replace("_", " ").upper()
        print(f"\n  {icon} Validation: {label}  (score {score}/100)")
        print(f"  {v.get('validation_summary','')}")
        for i in issues:
            print(f"  ISSUE [{i.get('type','')}]: {i.get('title','')} — Fix: {i.get('recommendation','')}")
        for w in warnings:
            print(f"  WARNING [{w.get('type','')}]: {w.get('title','')} — {w.get('description','')}")
        return

    label = decision.replace("_", " ").upper()
    console.print(Panel(
        f"[bold {color}]{icon}  {label}[/bold {color}]  ·  Quality score: {score}/100\n"
        f"[dim]{v.get('validation_summary', '')}[/dim]",
        title="[bold]Agent 2 — Scope Validation[/bold]",
        border_style=color, padding=(0, 1),
    ))

    if issues:
        t = Table("Type", "Title", "Recommendation", box=None, padding=(0, 2))
        for i in issues:
            t.add_row(
                f"[red]{i.get('type','')}[/red]",
                f"[bold]{i.get('title','')}[/bold]",
                i.get("recommendation", "")
            )
        console.print(t)

    if warnings:
        t = Table("Type", "Title", "Description", box=None, padding=(0, 2))
        for w in warnings:
            t.add_row(
                f"[yellow]{w.get('type','')}[/yellow]",
                f"[bold]{w.get('title','')}[/bold]",
                w.get("description", "")
            )
        console.print(t)

    approved = v.get("approved_items", [])
    if approved:
        if HAS_RICH:
            console.print(f"  [dim]Validated items: {' · '.join(approved)}[/dim]")

# ─── DOCX EXPORT ──────────────────────────────────────────────────────────────

def hex_to_rgb(hex_color: str):
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def export_docx(sow_markdown: str, output_path: Path, eng_type: str):
    """Convert SOW markdown to a formatted Word document (charcoal & gold theme)."""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    charcoal_rgb = hex_to_rgb("2C2C2C")
    gold_rgb     = hex_to_rgb("C9A84C")
    white_rgb    = (255, 255, 255)
    gold_light   = "F5E9C8"

    lines = sow_markdown.split("\n")

    def add_heading(text: str, level: int):
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text.strip())
        p    = doc.add_heading(text, level=min(level, 3))
        run  = p.runs[0] if p.runs else p.add_run(text)
        if level == 1:
            run.font.color.rgb = RGBColor(*charcoal_rgb)
            run.font.size      = Pt(20)
            run.font.bold      = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            # Gold bottom border
            pPr    = p._p.get_or_add_pPr()
            pBdr   = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "12")
            bottom.set(qn("w:space"), "4")
            bottom.set(qn("w:color"), "C9A84C")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif level == 2:
            run.font.color.rgb = RGBColor(*white_rgb)
            run.font.size      = Pt(11)
            run.font.bold      = True
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "2C2C2C")
            pPr.append(shd)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
        elif level == 3:
            run.font.color.rgb = RGBColor(*charcoal_rgb)
            run.font.size      = Pt(12)
            run.font.bold      = True

    def add_paragraph(text: str):
        if not text.strip():
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        parts = re.split(r'\*\*(.*?)\*\*', text.strip())
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            run.font.name      = "Calibri"
            run.font.size      = Pt(11)
            run.font.color.rgb = RGBColor(*charcoal_rgb)
            if idx % 2 == 1:
                run.font.bold = True

    def add_table_md(md_rows: list):
        rows = [
            r for r in md_rows
            if r.strip().startswith("|")
            and not re.match(r"^\|[-| :]+\|$", r.strip())
        ]
        if len(rows) < 2:
            return
        data = []
        for row in rows:
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            data.append(cells)
        if not data:
            return
        max_cols = max(len(r) for r in data)
        t = doc.add_table(rows=len(data), cols=max_cols)
        t.style     = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for r_idx, row in enumerate(data):
            for c_idx, cell_text in enumerate(row[:max_cols]):
                cell  = t.cell(r_idx, c_idx)
                clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text.strip())
                p     = cell.paragraphs[0]
                run   = p.add_run(clean)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                is_total = clean.lower().startswith("total")
                if r_idx == 0:
                    run.font.bold      = True
                    run.font.color.rgb = RGBColor(*white_rgb)
                    set_cell_bg(cell, "2C2C2C")
                elif is_total:
                    run.font.bold      = True
                    run.font.color.rgb = RGBColor(*charcoal_rgb)
                    set_cell_bg(cell, gold_light)
                elif r_idx % 2 == 0:
                    run.font.color.rgb = RGBColor(*charcoal_rgb)
                    set_cell_bg(cell, "F8FAFC")
                else:
                    run.font.color.rgb = RGBColor(*charcoal_rgb)
        doc.add_paragraph()

    # ── Parse and render markdown ─────────────────────────────────────────────
    table_buf = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# ") and not line.startswith("## "):
            if table_buf: add_table_md(table_buf); table_buf = []
            add_heading(line[2:], 1)
        elif line.startswith("## "):
            if table_buf: add_table_md(table_buf); table_buf = []
            add_heading(line[3:], 2)
        elif line.startswith("### "):
            if table_buf: add_table_md(table_buf); table_buf = []
            add_heading(line[4:], 3)
        elif line.strip().startswith("|"):
            table_buf.append(line)
        elif re.match(r"^[-─]+$", line.strip()):
            if table_buf: add_table_md(table_buf); table_buf = []
        elif line.strip():
            if table_buf: add_table_md(table_buf); table_buf = []
            if not line.strip().startswith("*This Statement"):
                add_paragraph(line)
        else:
            if table_buf: add_table_md(table_buf); table_buf = []
        i += 1

    if table_buf:
        add_table_md(table_buf)

    doc.save(str(output_path))

# ─── INTERACTIVE HELPERS ──────────────────────────────────────────────────────

def select_from_list(prompt: str, options: list, default: int = 0) -> str:
    if HAS_RICH:
        console.print(f"\n[bold]{prompt}[/bold]")
        for i, opt in enumerate(options, 1):
            marker = "[#C9A84C]→[/#C9A84C]" if i - 1 == default else " "
            console.print(f"  {marker} [{i}] {opt}")
        choice = console.input(f"  [dim]Enter number (default {default+1}): [/dim]").strip()
    else:
        print(f"\n{prompt}")
        for i, opt in enumerate(options, 1):
            print(f"  [{i}] {opt}")
        choice = input(f"  Enter number (default {default+1}): ").strip()

    if not choice:
        return options[default]
    try:
        return options[max(0, min(int(choice) - 1, len(options) - 1))]
    except ValueError:
        return options[default]

# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="SOW Orchestrator — Autonomous 3-agent SOW generation pipeline",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python sow_generator.py notes.txt
  python sow_generator.py notes.txt --provider openai
  python sow_generator.py notes.txt --type "Telco NFV Transformation" --industry "Telco / CSP"
  python sow_generator.py notes.txt --no-docx --output my_sow.md
  python sow_generator.py --interactive

Environment variables:
  ANTHROPIC_API_KEY   — Anthropic API key (Claude)
  OPENAI_API_KEY      — OpenAI API key

AI Portfolio · Daniel Mazzini · github.com/danvzla
        """
    )
    parser.add_argument("notes_file",    nargs="?",          help="Path to discovery notes text file")
    parser.add_argument("--provider",    default="claude",   choices=["claude", "openai"],
                        help="AI provider: claude (default) or openai")
    parser.add_argument("--type",        default=None,       help="Engagement type")
    parser.add_argument("--size",        default=None,       help="Engagement size")
    parser.add_argument("--industry",    default=None,       help="Customer industry")
    parser.add_argument("--output",      default=None,       help="Output file path (without extension)")
    parser.add_argument("--no-docx",     action="store_true",help="Skip Word document export")
    parser.add_argument("--interactive", action="store_true",help="Interactive mode — prompts for all inputs")
    parser.add_argument("--api-key",     default=None,       help="API key (Claude or OpenAI)")
    args = parser.parse_args()

    provider = args.provider

    # ── Provider setup ────────────────────────────────────────────────────────
    claude_client = None
    openai_client = None

    if provider == "claude":
        if not HAS_ANTHROPIC:
            print("\n[ERROR] anthropic package not found. Run: pip install anthropic\n")
            sys.exit(1)
        api_key = args.api_key or os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            api_key = (console.input("[bold]Anthropic API key:[/bold] ")
                       if HAS_RICH else input("Anthropic API key: ")).strip()
        if not api_key:
            print("Error: No Anthropic API key provided.")
            sys.exit(1)
        claude_client = anthropic.Anthropic(api_key=api_key)

    elif provider == "openai":
        if not HAS_OPENAI:
            print("\n[ERROR] openai package not found. Run: pip install openai\n")
            sys.exit(1)
        api_key = args.api_key or os.environ.get("OPENAI_API_KEY", "")
        if not api_key:
            api_key = (console.input("[bold]OpenAI API key:[/bold] ")
                       if HAS_RICH else input("OpenAI API key: ")).strip()
        if not api_key:
            print("Error: No OpenAI API key provided.")
            sys.exit(1)
        openai_client = openai_lib.OpenAI(api_key=api_key)

    print_header(provider)

    # ── Discovery notes ───────────────────────────────────────────────────────
    if args.notes_file:
        path = Path(args.notes_file)
        if not path.exists():
            print(f"Error: File not found: {path}")
            sys.exit(1)
        notes = path.read_text(encoding="utf-8").strip()
        msg = f"Notes loaded: {path} ({len(notes)} chars)"
        console.print(f"[dim]{msg}[/dim]") if HAS_RICH else print(msg)
    else:
        if HAS_RICH:
            console.print("[bold]Paste your discovery notes below.[/bold]")
            console.print("[dim]Type END on a new line when done:[/dim]\n")
        else:
            print("\nPaste discovery notes. Type END when done:\n")
        lines = []
        while True:
            try:
                line = input()
                if line.strip().upper() == "END":
                    break
                lines.append(line)
            except EOFError:
                break
        notes = "\n".join(lines).strip()
        if not notes:
            print("Error: No discovery notes provided.")
            sys.exit(1)

    # ── Engagement context ────────────────────────────────────────────────────
    eng_type = (args.type     or select_from_list("Engagement type:",  ENGAGEMENT_TYPES))
    eng_size = (args.size     or select_from_list("Engagement size:",  ENGAGEMENT_SIZES, default=1))
    industry = (args.industry or select_from_list("Customer industry:", INDUSTRIES))

    if HAS_RICH:
        console.print(f"\n[dim]Engagement:[/dim] {eng_type}  ·  {eng_size}  ·  {industry}\n")
    else:
        print(f"\nEngagement: {eng_type}  |  {eng_size}  |  {industry}\n")

    # ── Output paths ──────────────────────────────────────────────────────────
    timestamp   = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_type   = re.sub(r"[^a-z0-9]+", "_", eng_type.lower())[:30]
    output_stem = args.output or f"SOW_{safe_type}_{timestamp}"
    output_md   = Path(output_stem + ".md")
    output_docx = Path(output_stem + ".docx")

    # ─────────────────────────────────────────────────────────────────────────
    # AGENT 1 — DISCOVERY ANALYST
    # ─────────────────────────────────────────────────────────────────────────
    print_agent_start(1, "Discovery Analyst",
                      "Reading notes · structuring scope · mapping deliverables and risks")
    try:
        if HAS_RICH:
            with Progress(SpinnerColumn(), TextColumn("[dim]{task.description}[/dim]"),
                          TimeElapsedColumn(), console=console) as prog:
                prog.add_task("Analyzing discovery notes...", total=None)
                brief = run_agent1(provider, notes, eng_type, eng_size, industry,
                                   claude_client, openai_client)
        else:
            print("  Running Agent 1...")
            brief = run_agent1(provider, notes, eng_type, eng_size, industry,
                               claude_client, openai_client)
    except Exception as e:
        print_agent_error(1, str(e))
        sys.exit(1)

    summary = (f"Brief complete · {len(brief.get('in_scope',[]))} workstreams · "
               f"{len(brief.get('deliverables',[]))} deliverables · "
               f"{len(brief.get('risks',[]))} risks · "
               f"duration: {brief.get('estimated_duration','—')}")
    print_agent_done(1, summary)
    print_brief_summary(brief)

    # ─────────────────────────────────────────────────────────────────────────
    # AGENT 2 — SCOPE VALIDATOR
    # ─────────────────────────────────────────────────────────────────────────
    print_agent_start(2, "Scope Validator",
                      "Checking for contradictions · gaps · ambiguities · autonomous decision")
    try:
        if HAS_RICH:
            with Progress(SpinnerColumn(), TextColumn("[dim]{task.description}[/dim]"),
                          TimeElapsedColumn(), console=console) as prog:
                prog.add_task("Validating scope brief...", total=None)
                validation = run_agent2(provider, brief, eng_type, eng_size, industry,
                                        claude_client, openai_client)
        else:
            print("  Running Agent 2...")
            validation = run_agent2(provider, brief, eng_type, eng_size, industry,
                                    claude_client, openai_client)
    except Exception as e:
        print_agent_error(2, str(e))
        sys.exit(1)

    decision = validation.get("decision", "approved")
    print_validation_result(validation)

    if decision == "rejected":
        print_agent_error(2, "Scope rejected. Fix the issues above and rerun.")
        if HAS_RICH:
            console.print("\n[red]Agent 3 blocked.[/red] Edit your discovery notes and run again.\n")
        else:
            print("\nAgent 3 blocked. Edit notes and rerun.")
        sys.exit(2)

    print_agent_done(2,
        f"Decision: {decision.upper()} · Quality score: {validation.get('quality_score', 0)}/100")

    # ─────────────────────────────────────────────────────────────────────────
    # AGENT 3 — SOW DRAFTER
    # ─────────────────────────────────────────────────────────────────────────
    print_agent_start(3, "SOW Drafter",
                      "Writing 14-section Statement of Work with tables and effort estimates")
    try:
        if HAS_RICH:
            with Progress(SpinnerColumn(), TextColumn("[dim]{task.description}[/dim]"),
                          TimeElapsedColumn(), console=console) as prog:
                prog.add_task("Drafting Statement of Work...", total=None)
                sow_markdown = run_agent3(provider, brief, validation,
                                          eng_type, eng_size, industry,
                                          claude_client, openai_client)
        else:
            print("  Running Agent 3...")
            sow_markdown = run_agent3(provider, brief, validation,
                                      eng_type, eng_size, industry,
                                      claude_client, openai_client)
    except Exception as e:
        print_agent_error(3, str(e))
        sys.exit(1)

    # Save Markdown
    output_md.write_text(sow_markdown, encoding="utf-8")
    line_count = len(sow_markdown.splitlines())
    print_agent_done(3, f"SOW drafted · {line_count} lines · saved → {output_md}")

    # Save DOCX
    if not args.no_docx:
        if HAS_DOCX:
            try:
                export_docx(sow_markdown, output_docx, eng_type)
                msg = f"Word document saved → {output_docx}"
                console.print(f"  [green]✓[/green]  {msg}") if HAS_RICH else print(f"  ✓ {msg}")
            except Exception as e:
                msg = f"DOCX export failed: {e}  (Markdown saved successfully)"
                console.print(f"  [yellow]⚠[/yellow]  {msg}") if HAS_RICH else print(f"  ⚠ {msg}")
        else:
            msg = "python-docx not installed — skipping Word export. Run: pip install python-docx"
            console.print(f"  [yellow]⚠[/yellow]  {msg}") if HAS_RICH else print(f"  ⚠ {msg}")

    # ── Summary ───────────────────────────────────────────────────────────────
    provider_label = (f"Claude ({CLAUDE_MODEL})" if provider == "claude"
                      else f"OpenAI ({OPENAI_MODEL})")
    if HAS_RICH:
        console.print()
        console.print(Panel(
            f"[bold green]✓ SOW generation complete[/bold green]\n\n"
            f"[dim]Markdown:[/dim]   {output_md}\n"
            f"[dim]Word doc:[/dim]   {output_docx if (not args.no_docx and HAS_DOCX) else 'skipped'}\n\n"
            f"[dim]Provider:[/dim]   {provider_label}\n"
            f"[dim]Quality:[/dim]    {validation.get('quality_score', 0)}/100  ·  "
            f"[dim]Validation:[/dim] {decision.upper()}  ·  "
            f"[dim]Duration:[/dim]   {brief.get('estimated_duration', '—')}",
            border_style="green", padding=(1, 2),
        ))
        console.print()
        console.print("[dim]AI SOW Orchestrator · Daniel Mazzini · github.com/danvzla[/dim]\n")
    else:
        print(f"\n✓ Complete.")
        print(f"  Markdown:  {output_md}")
        if not args.no_docx and HAS_DOCX:
            print(f"  Word doc:  {output_docx}")
        print(f"  Provider:  {provider_label}")
        print(f"  Score:     {validation.get('quality_score', 0)}/100  |  {decision.upper()}")


if __name__ == "__main__":
    main()
